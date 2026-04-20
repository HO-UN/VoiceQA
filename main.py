import sys
import os
import threading
import queue
import time
from datetime import datetime

# GUI库
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QTextEdit, QPushButton, QLabel, QComboBox, QLineEdit, 
    QFileDialog, QTabWidget, QGroupBox, QFormLayout, QCheckBox
)
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QFont

# 音频处理库
import sounddevice as sd
import numpy as np
import wave
import pyaudiowpatch as pyaudio

# ASR库
import vosk
import json

# OpenAI API
import openai

# 嵌入和检索
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.neighbors import NearestNeighbors
import json

from PyQt5.QtCore import pyqtSignal

class VoiceQA(QMainWindow):
    # 定义类级别的信号
    update_transcript_signal = pyqtSignal(str)
    update_answer_signal = pyqtSignal(str)
    update_status_signal = pyqtSignal(str)
    
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Voice QA Assistant")
        self.setGeometry(100, 100, 800, 600)
        
        # 初始化配置
        self.config = {
            "audio_source": "",
            "audio_source_type": "input",  # input or output
            "openai_api_key": "sk-vCarGAu1wVOSlDQ4C54c3f88A69e41779893Ad6a81A76619",
            "openai_base_url":
             "https://aihubmix.com/v1",
            "openai_model": "coding-minimax-m2.7-free",
            "asr_model": "vosk-model-small-cn-0.22",
            "knowledge_base_folder": "knowledge_base"
        }
        
        # 加载配置
        self.load_config()
        
        # 初始化ASR模型
        self.asr_model = None
        self.vectorizer = TfidfVectorizer(analyzer='char', ngram_range=(1, 2))
        self.knn = None
        self.knowledge_chunks = []
        
        # 音频队列
        self.audio_queue = queue.Queue()
        self.is_recording = False
        self.audio_thread = None
        
        # 初始化UI
        self.init_ui()
        
        # 连接信号
        self.update_transcript_signal.connect(self.update_transcript)
        self.update_answer_signal.connect(self.update_answer)
        self.update_status_signal.connect(self.update_status)
        
        # 加载模型
        self.load_models()
        self.load_knowledge_base()
    
    def init_ui(self):
        # 主布局
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # 选项卡
        self.tabs = QTabWidget()
        main_layout.addWidget(self.tabs)
        
        # 主界面
        main_tab = QWidget()
        main_layout_tab = QVBoxLayout(main_tab)
        
        # 音频源选择
        audio_source_group = QGroupBox("音频源")
        audio_source_layout = QVBoxLayout()
        
        # 音频源类型选择
        source_type_layout = QHBoxLayout()
        source_type_layout.addWidget(QLabel("音频源类型:"))
        self.source_type_combo = QComboBox()
        self.source_type_combo.addItems(["输入设备", "输出设备"])
        self.source_type_combo.currentIndexChanged.connect(self.refresh_audio_sources)
        source_type_layout.addWidget(self.source_type_combo)
        audio_source_layout.addLayout(source_type_layout)
        
        # 音频设备选择
        device_layout = QHBoxLayout()
        self.audio_source_combo = QComboBox()
        self.refresh_audio_sources()
        device_layout.addWidget(QLabel("选择设备:"))
        device_layout.addWidget(self.audio_source_combo)
        refresh_button = QPushButton("刷新")
        refresh_button.clicked.connect(self.refresh_audio_sources)
        device_layout.addWidget(refresh_button)
        audio_source_layout.addLayout(device_layout)
        
        audio_source_group.setLayout(audio_source_layout)
        main_layout_tab.addWidget(audio_source_group)
        
        # 控制按钮
        control_layout = QHBoxLayout()
        self.start_button = QPushButton("启动自动采集")
        self.start_button.clicked.connect(self.start_recording)
        self.stop_button = QPushButton("停止采集")
        self.stop_button.clicked.connect(self.stop_recording)
        self.stop_button.setEnabled(False)
        control_layout.addWidget(self.start_button)
        control_layout.addWidget(self.stop_button)
        main_layout_tab.addLayout(control_layout)
        
        # 转录结果
        transcript_group = QGroupBox("音频转文字")
        transcript_layout = QVBoxLayout()
        self.transcript_text = QTextEdit()
        self.transcript_text.setReadOnly(True)
        transcript_layout.addWidget(self.transcript_text)
        transcript_group.setLayout(transcript_layout)
        main_layout_tab.addWidget(transcript_group)
        
        # AI回答
        answer_group = QGroupBox("AI回答")
        answer_layout = QVBoxLayout()
        self.answer_text = QTextEdit()
        self.answer_text.setReadOnly(True)
        answer_layout.addWidget(self.answer_text)
        answer_group.setLayout(answer_layout)
        main_layout_tab.addWidget(answer_group)
        
        self.tabs.addTab(main_tab, "主界面")
        
        # 配置界面
        config_tab = QWidget()
        config_layout = QVBoxLayout(config_tab)
        
        # OpenAI配置
        openai_group = QGroupBox("OpenAI配置")
        openai_layout = QFormLayout()
        self.api_key_input = QLineEdit()
        self.api_key_input.setText(self.config["openai_api_key"])
        openai_layout.addRow("API Key:", self.api_key_input)
        self.base_url_input = QLineEdit()
        self.base_url_input.setText(self.config["openai_base_url"])
        openai_layout.addRow("Base URL:", self.base_url_input)
        self.openai_model_input = QLineEdit()
        self.openai_model_input.setText(self.config["openai_model"])
        openai_layout.addRow("模型名称:", self.openai_model_input)
        openai_layout.addRow(QLabel("提示: 可输入任意OpenAI模型名称，如gpt-3.5-turbo-instruct、gpt-4等"))
        openai_group.setLayout(openai_layout)
        config_layout.addWidget(openai_group)
        
        # 模型配置
        model_group = QGroupBox("模型配置")
        model_layout = QFormLayout()
        self.asr_model_input = QLineEdit()
        self.asr_model_input.setText(self.config["asr_model"])
        model_layout.addRow("ASR模型名称:", self.asr_model_input)
        model_layout.addRow(QLabel("提示: 模型名称需与Vosk模型文件名一致，如vosk-model-small-cn-0.22"))
        model_group.setLayout(model_layout)
        config_layout.addWidget(model_group)
        
        # 知识库配置
        kb_group = QGroupBox("知识库配置")
        kb_layout = QHBoxLayout()
        self.kb_path_input = QLineEdit()
        self.kb_path_input.setText(self.config["knowledge_base_folder"])
        kb_layout.addWidget(self.kb_path_input)
        browse_button = QPushButton("浏览")
        browse_button.clicked.connect(self.browse_kb)
        kb_layout.addWidget(browse_button)
        kb_group.setLayout(kb_layout)
        config_layout.addWidget(kb_group)
        
        # 添加提示信息
        config_layout.addWidget(QLabel("提示: 知识库文件夹中可包含PDF、Word文档，启动时会自动加载到内存"))
        
        # 保存配置按钮
        save_button = QPushButton("保存配置")
        save_button.clicked.connect(self.save_config)
        config_layout.addWidget(save_button)
        
        self.tabs.addTab(config_tab, "配置")
        
        # 状态栏
        self.statusBar().showMessage("就绪")
    
    def refresh_audio_sources(self):
        self.audio_source_combo.clear()
        
        # 使用pyaudiowpatch获取所有音频设备
        p = pyaudio.PyAudio()
        info = p.get_host_api_info_by_index(0)
        numdevices = info.get('deviceCount')
        
        # 获取当前选择的音频源类型
        source_type = self.source_type_combo.currentText()
        
        for i in range(0, numdevices):
            device_info = p.get_device_info_by_host_api_device_index(0, i)
            if source_type == "输入设备":
                if device_info.get('maxInputChannels') > 0:
                    name = device_info.get('name')
                    self.audio_source_combo.addItem(name)
            else:  # 输出设备
                if device_info.get('maxOutputChannels') > 0:
                    name = device_info.get('name')
                    self.audio_source_combo.addItem(name)
        
        p.terminate()
    
    def browse_kb(self):
        folder_path = QFileDialog.getExistingDirectory(self, "选择知识库文件夹")
        if folder_path:
            self.kb_path_input.setText(folder_path)
    
    def load_config(self):
        if os.path.exists("config.json"):
            with open("config.json", "r", encoding="utf-8") as f:
                self.config.update(json.load(f))
    
    def save_config(self):
        self.config["audio_source"] = self.audio_source_combo.currentText()
        self.config["audio_source_type"] = "input" if self.source_type_combo.currentText() == "输入设备" else "output"
        self.config["openai_api_key"] = self.api_key_input.text()
        self.config["openai_base_url"] = self.base_url_input.text()
        self.config["openai_model"] = self.openai_model_input.text()
        self.config["asr_model"] = self.asr_model_input.text()
        self.config["knowledge_base_folder"] = self.kb_path_input.text()
        
        with open("config.json", "w", encoding="utf-8") as f:
            json.dump(self.config, f, indent=2, ensure_ascii=False)
        
        self.statusBar().showMessage("配置已保存")
        
        # 重新加载模型
        self.load_models()
        self.load_knowledge_base()
    
    def load_models(self):
        try:
            # 加载ASR模型
            self.update_status_signal.emit("加载ASR模型...")
            # 下载并使用Vosk模型
            import os
            import requests
            import zipfile
            
            # 确保asr_model配置存在
            if "asr_model" not in self.config:
                self.config["asr_model"] = "vosk-model-small-cn-0.22"
                print("未找到asr_model配置，使用默认值: vosk-model-small-cn-0.22")
            
            model_name = self.config["asr_model"]
            model_path = os.path.abspath(model_name)
            
            print(f"尝试加载ASR模型: {model_name}")
            print(f"模型路径: {model_path}")
            print(f"当前工作目录: {os.getcwd()}")
            
            # 检查模型路径是否存在
            if not os.path.exists(model_path):
                print(f"模型路径不存在: {model_path}")
                # 尝试使用相对路径
                model_path = os.path.join(os.getcwd(), model_name)
                print(f"尝试使用相对路径: {model_path}")
                
                if not os.path.exists(model_path):
                    # 下载模型
                    url = f"https://alphacephei.com/vosk/models/{model_name}.zip"
                    zip_path = f"{model_name}.zip"
                    
                    print(f"模型不存在，开始下载: {url}")
                    self.update_status_signal.emit("下载ASR模型...")
                    
                    try:
                        print(f"开始下载模型文件: {zip_path}")
                        response = requests.get(url, stream=True, timeout=60)
                        
                        # 检查响应状态
                        if response.status_code != 200:
                            print(f"下载失败，状态码: {response.status_code}")
                            self.update_status_signal.emit(f"下载ASR模型失败: 状态码 {response.status_code}")
                            return
                        
                        total_size = int(response.headers.get('content-length', 0))
                        downloaded_size = 0
                        
                        with open(zip_path, 'wb') as f:
                            for chunk in response.iter_content(chunk_size=8192):
                                if chunk:
                                    f.write(chunk)
                                    downloaded_size += len(chunk)
                                    if total_size > 0:
                                        progress = int((downloaded_size / total_size) * 100)
                                        print(f"下载进度: {progress}%")
                        
                        print(f"下载完成，文件大小: {os.path.getsize(zip_path)} bytes")
                        
                        # 解压模型
                        print("开始解压模型...")
                        self.update_status_signal.emit("解压ASR模型...")
                        
                        try:
                            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                                zip_ref.extractall('.')
                            print("模型解压完成")
                        except zipfile.BadZipFile:
                            print("解压失败: 无效的ZIP文件")
                            self.update_status_signal.emit("解压ASR模型失败: 无效的ZIP文件")
                            # 删除损坏的ZIP文件
                            if os.path.exists(zip_path):
                                os.remove(zip_path)
                            return
                        
                        # 删除zip文件
                        if os.path.exists(zip_path):
                            os.remove(zip_path)
                            print("已删除ZIP文件")
                    except Exception as e:
                        print(f"下载或解压模型失败: {str(e)}")
                        import traceback
                        traceback.print_exc()
                        self.update_status_signal.emit(f"下载ASR模型失败: {str(e)}")
                        return
            
            # 确认模型路径存在
            if not os.path.exists(model_path):
                print(f"模型路径仍然不存在: {model_path}")
                self.update_status_signal.emit("ASR模型路径不存在")
                return
            
            # 检查模型路径是否是目录
            if not os.path.isdir(model_path):
                print(f"模型路径不是目录: {model_path}")
                self.update_status_signal.emit("ASR模型路径不是目录")
                return
            
            # 检查模型目录中的文件
            model_files = os.listdir(model_path)
            print(f"模型目录中的文件: {model_files}")
            
            # 检查必要的模型文件和目录
            required_items = ["ivector", "graph", "am"]
            missing_items = []
            for item in required_items:
                if item not in model_files:
                    missing_items.append(item)
            
            if missing_items:
                print(f"缺少必要的模型文件或目录: {missing_items}")
                self.update_status_signal.emit(f"ASR模型缺少必要文件或目录: {missing_items}")
                return
            
            # 检查am目录中是否有模型文件
            am_dir = os.path.join(model_path, "am")
            if os.path.exists(am_dir):
                am_files = os.listdir(am_dir)
                print(f"am目录中的文件: {am_files}")
                # 检查是否有final.mdl文件
                if "final.mdl" not in am_files:
                    print("缺少final.mdl模型文件")
                    self.update_status_signal.emit("ASR模型缺少final.mdl文件")
                    return
            else:
                print("缺少am目录")
                self.update_status_signal.emit("ASR模型缺少am目录")
                return
            
            try:
                print(f"加载模型: {model_path}")
                self.asr_model = vosk.Model(model_path)
                print("ASR模型加载成功")
                self.update_status_signal.emit("ASR模型加载成功")
            except Exception as e:
                print(f"加载ASR模型失败: {str(e)}")
                import traceback
                traceback.print_exc()
                self.update_status_signal.emit(f"加载ASR模型失败: {str(e)}")
                return
            
            # 配置OpenAI
            openai.api_key = self.config["openai_api_key"]
            if self.config["openai_base_url"]:
                openai.api_base = self.config["openai_base_url"]
            
            self.update_status_signal.emit("模型加载完成")
            print("所有模型加载完成")
        except Exception as e:
            error_msg = f"模型加载失败: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            self.update_status_signal.emit(error_msg)
    
    def extract_text_from_pdf(self, pdf_path):
        """从PDF文件中提取文本"""
        import PyPDF2
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"提取PDF文本失败: {e}")
        return text
    
    def extract_text_from_docx(self, docx_path):
        """从Word文档中提取文本"""
        text = ""
        try:
            print(f"尝试从Word文档提取文本: {docx_path}")
            # 确保文件存在
            if not os.path.exists(docx_path):
                print(f"文件不存在: {docx_path}")
                return text
            
            # 导入Document类
            from docx import Document
            
            # 打开文档
            print(f"打开文档: {docx_path}")
            doc = Document(docx_path)
            
            # 提取段落文本
            print(f"提取段落文本，共 {len(doc.paragraphs)} 个段落")
            for i, paragraph in enumerate(doc.paragraphs):
                if i < 5:  # 只打印前5个段落的前50个字符
                    preview = paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text
                    print(f"段落 {i+1}: {preview}")
                text += paragraph.text + "\n"
            
            print(f"成功提取Word文本，总长度: {len(text)}")
        except Exception as e:
            error_msg = f"提取Word文本失败: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
        return text
    
    def split_text_into_chunks(self, text, chunk_size=1000, overlap=200):
        """将文本分割成块"""
        chunks = []
        start = 0
        text_length = len(text)
        print(f"开始分割文本，长度: {text_length}, 块大小: {chunk_size}, 重叠: {overlap}")
        
        # 确保chunk_size大于overlap，防止无限循环
        if chunk_size <= overlap:
            overlap = chunk_size // 2
            print(f"调整重叠大小为: {overlap}")
        
        try:
            # 限制最大块数，防止内存溢出
            max_chunks = 1000
            chunk_count = 0
            
            while start < text_length and chunk_count < max_chunks:
                end = min(start + chunk_size, text_length)
                chunk = text[start:end]
                chunks.append(chunk)
                chunk_count += 1
                print(f"添加块 {chunk_count}，长度: {len(chunk)}")
                
                # 计算下一个起始位置
                next_start = end - overlap
                
                # 防止无限循环
                if next_start <= start or next_start >= text_length:
                    break
                
                start = next_start
                
                # 防止内存溢出
                if chunk_count >= max_chunks:
                    print("达到最大块数限制，停止分割")
                    break
        except Exception as e:
            print(f"分割文本时出错: {str(e)}")
            import traceback
            traceback.print_exc()
            # 返回至少一个块
            if not chunks and text:
                chunks.append(text[:min(chunk_size, text_length)])
        
        print(f"分割完成，共 {len(chunks)} 个块")
        return chunks
    
    def update_transcript(self, text):
        """通过信号更新转录文本"""
        if self.transcript_text:
            self.transcript_text.append(text)
    
    def update_answer(self, text):
        """通过信号更新回答文本"""
        if self.answer_text:
            self.answer_text.append(text)
    
    def update_status(self, text):
        """通过信号更新状态栏"""
        self.statusBar().showMessage(text)
    
    def load_knowledge_base(self):
        try:
            # 确保knowledge_base_folder配置存在
            if "knowledge_base_folder" not in self.config:
                self.config["knowledge_base_folder"] = "knowledge_base"
                
            knowledge_base_folder = self.config["knowledge_base_folder"]
            print(f"尝试加载知识库文件夹: {knowledge_base_folder}")
            
            # 确保文件夹存在
            if not os.path.exists(knowledge_base_folder):
                # 创建默认的知识库文件夹
                try:
                    os.makedirs(knowledge_base_folder, exist_ok=True)
                    print(f"创建了知识库文件夹: {knowledge_base_folder}")
                except Exception as e:
                    error_msg = f"无法创建知识库文件夹: {str(e)}"
                    print(error_msg)
                    self.update_status_signal.emit(error_msg)
                    return
            
            # 确保路径是文件夹
            if not os.path.isdir(knowledge_base_folder):
                error_msg = "知识库路径不是文件夹"
                print(error_msg)
                self.update_status_signal.emit(error_msg)
                return
            
            # 初始化知识库
            self.knowledge_chunks = []
            
            try:
                # 遍历文件夹中的所有文件
                print("开始遍历文件夹中的文件...")
                for root, dirs, files in os.walk(knowledge_base_folder):
                    print(f"在目录 {root} 中找到 {len(files)} 个文件")
                    for file in files:
                        file_path = os.path.join(root, file)
                        file_ext = os.path.splitext(file)[1].lower()
                        print(f"处理文件: {file_path}, 扩展名: {file_ext}")
                        
                        # 根据文件类型提取文本
                        try:
                            if file_ext == '.pdf':
                                print(f"提取PDF文本: {file_path}")
                                text = self.extract_text_from_pdf(file_path)
                            elif file_ext == '.docx':
                                print(f"提取Word文本: {file_path}")
                                text = self.extract_text_from_docx(file_path)
                            else:
                                print(f"跳过不支持的文件类型: {file_ext}")
                                continue  # 跳过不支持的文件类型
                            
                            # 将文本分割成块
                            print(f"分割文本为块，文本长度: {len(text)}")
                            chunks = self.split_text_into_chunks(text)
                            self.knowledge_chunks.extend(chunks)
                            print(f"添加了 {len(chunks)} 个文本块")
                        except Exception as e:
                            print(f"处理文件 {file_path} 时出错: {str(e)}")
                            import traceback
                            traceback.print_exc()
                            continue
            except Exception as e:
                print(f"遍历文件夹时出错: {str(e)}")
                import traceback
                traceback.print_exc()
            
            print(f"知识库文本块总数: {len(self.knowledge_chunks)}")
            
            # 使用简单的向量存储方案
            try:
                if self.knowledge_chunks:
                    # 训练TF-IDF向量化器
                    print("训练TF-IDF向量化器...")
                    self.vectorizer.fit(self.knowledge_chunks)
                    
                    # 转换文本为向量
                    print("转换文本为向量...")
                    vectors = self.vectorizer.transform(self.knowledge_chunks)
                    
                    # 保存向量和文本块到文件
                    print("保存向量到文件...")
                    import pickle
                    with open("vector_store.pkl", "wb") as f:
                        pickle.dump({
                            "chunks": self.knowledge_chunks,
                            "vectors": vectors
                        }, f)
                    
                    # 训练KNN模型
                    print("训练KNN模型...")
                    self.knn = NearestNeighbors(n_neighbors=3, metric="cosine")
                    self.knn.fit(vectors)
                    
                    success_msg = f"知识库加载完成，共 {len(self.knowledge_chunks)} 个文本块"
                    print(success_msg)
                    self.update_status_signal.emit(success_msg)
                else:
                    # 检查是否有之前保存的向量
                    if os.path.exists("vector_store.pkl"):
                        print("加载之前保存的向量...")
                        import pickle
                        with open("vector_store.pkl", "rb") as f:
                            data = pickle.load(f)
                            self.knowledge_chunks = data["chunks"]
                            vectors = data["vectors"]
                            
                            # 训练KNN模型
                            self.knn = NearestNeighbors(n_neighbors=3, metric="cosine")
                            self.knn.fit(vectors)
                            
                            success_msg = f"加载之前的知识库，共 {len(self.knowledge_chunks)} 个文本块"
                            print(success_msg)
                            self.update_status_signal.emit(success_msg)
                    else:
                        # 初始化空的向量器和KNN模型
                        print("初始化空的向量器和KNN模型...")
                        try:
                            # 使用一个非空的默认文本进行初始化
                            default_text = "默认文本 初始化 向量器"
                            self.vectorizer.fit([default_text])
                            self.knn = NearestNeighbors(n_neighbors=3, metric="cosine")
                            self.knn.fit(self.vectorizer.transform([default_text]))
                            info_msg = "知识库为空，已初始化默认模型"
                            print(info_msg)
                            self.update_status_signal.emit(info_msg)
                        except Exception as e:
                            print(f"初始化默认模型时出错: {str(e)}")
                            # 即使初始化失败，也要继续运行
                            self.update_status_signal.emit("知识库模型初始化失败，但应用可以继续运行")
            except Exception as e:
                print(f"初始化模型时出错: {str(e)}")
                import traceback
                traceback.print_exc()
                # 即使模型初始化失败，也要继续运行
                self.update_status_signal.emit("知识库模型初始化失败，但应用可以继续运行")
                
        except Exception as e:
            error_msg = f"知识库加载失败: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            self.update_status_signal.emit(error_msg)
    
    def start_recording(self):
        self.is_recording = True
        self.start_button.setEnabled(False)
        self.stop_button.setEnabled(True)
        self.update_status_signal.emit("正在自动采集音频...")
        
        # 启动音频录制线程
        self.audio_thread = threading.Thread(target=self.record_audio)
        self.audio_thread.daemon = True
        self.audio_thread.start()
    
    def stop_recording(self):
        self.is_recording = False
        self.start_button.setEnabled(True)
        self.stop_button.setEnabled(False)
        self.update_status_signal.emit("自动采集已停止")
    
    def record_audio(self):
        # 获取选择的音频设备
        device_name = self.audio_source_combo.currentText()
        source_type = self.source_type_combo.currentText()
        
        # 停顿检测参数
        silence_threshold = 0.02  # 静默阈值，进一步降低以提高灵敏度
        silence_duration = 0.2  # 静默持续时间（秒），进一步缩短以更快触发转录
        silence_counter = 0
        audio_buffer = []
        frames_per_buffer = 1024
        samplerate = 16000
        
        print(f"停顿检测参数: 阈值={silence_threshold}, 持续时间={silence_duration}秒")
        print(f"音频设备: {device_name}, 类型: {source_type}")
        
        # 使用sounddevice录制音频
        def callback(indata, frames, time, status):
            nonlocal silence_counter, audio_buffer
            
            if status:
                print(f"音频状态: {status}")
            
            # 计算音频能量
            energy = np.linalg.norm(indata) / np.sqrt(len(indata))
            
            if energy > silence_threshold:
                # 有声音，重置静默计数器
                silence_counter = 0
                audio_buffer.append(indata.copy())
            else:
                # 静默，增加计数器
                silence_counter += frames / samplerate
                if silence_counter > silence_duration and len(audio_buffer) > 0:
                    # 检测到停顿，处理音频
                    print(f"检测到停顿，处理音频缓冲区，大小: {len(audio_buffer)} 帧")
                    self.process_audio_buffer(audio_buffer)
                    audio_buffer = []
                    silence_counter = 0
        
        try:
            if source_type == "输入设备":
                with sd.InputStream(callback=callback, channels=1, samplerate=samplerate, blocksize=frames_per_buffer):
                    while self.is_recording:
                        time.sleep(0.1)
            else:
                # 对于输出设备，使用环回录制
                with sd.InputStream(callback=callback, channels=1, samplerate=samplerate, blocksize=frames_per_buffer, device=device_name):
                    while self.is_recording:
                        time.sleep(0.1)
        except Exception as e:
            self.update_status_signal.emit(f"录音失败: {str(e)}")
    
    def process_audio_buffer(self, audio_buffer):
        # 检查音频缓冲区大小
        if not audio_buffer or len(audio_buffer) == 0:
            print("音频缓冲区为空，跳过处理")
            return
        
        # 合并音频数据
        audio = np.concatenate(audio_buffer, axis=0)
        print(f"处理音频数据，长度: {len(audio)}")
        
        # 检查音频能量，避免处理静音
        energy = np.linalg.norm(audio) / np.sqrt(len(audio))
        print(f"音频能量: {energy}")
        if energy < 0.001:
            print("音频能量过低，可能是静音，跳过处理")
            return
        
        # 保存为临时文件
        temp_file = f"temp_{datetime.now().timestamp()}.wav"
        print(f"保存音频到临时文件: {temp_file}")
        with wave.open(temp_file, 'wb') as wf:
            wf.setnchannels(1)
            wf.setsampwidth(2)
            wf.setframerate(16000)
            wf.writeframes((audio * 32767).astype(np.int16).tobytes())
        
        # 转录音频
        self.transcribe_audio(temp_file)
        
        # 删除临时文件
        if os.path.exists(temp_file):
            os.remove(temp_file)
            print(f"已删除临时文件: {temp_file}")
    
    def process_audio(self):
        while True:
            if not self.audio_queue.empty() and not self.is_recording:
                # 收集所有音频数据
                audio_data = []
                while not self.audio_queue.empty():
                    audio_data.append(self.audio_queue.get())
                
                if audio_data:
                    # 合并音频数据
                    audio = np.concatenate(audio_data, axis=0)
                    
                    # 保存为临时文件
                    temp_file = f"temp_{datetime.now().timestamp()}.wav"
                    with wave.open(temp_file, 'wb') as wf:
                        wf.setnchannels(1)
                        wf.setsampwidth(2)
                        wf.setframerate(16000)
                        wf.writeframes((audio * 32767).astype(np.int16).tobytes())
                    
                    # 转录音频
                    self.transcribe_audio(temp_file)
                    
                    # 删除临时文件
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
            
            time.sleep(0.1)
    
    def transcribe_audio(self, audio_file):
        try:
            print(f"开始转录音频文件: {audio_file}")
            self.update_status_signal.emit("正在转录音频...")
            
            # 检查ASR模型是否正确加载
            if not self.asr_model:
                error_msg = "ASR模型未加载"
                print(error_msg)
                self.update_status_signal.emit(error_msg)
                return
            
            # 使用Vosk转录
            recognizer = vosk.KaldiRecognizer(self.asr_model, 16000)
            
            # 检查音频文件是否存在
            if not os.path.exists(audio_file):
                error_msg = "音频文件不存在"
                print(error_msg)
                self.update_status_signal.emit(error_msg)
                return
            
            try:
                with wave.open(audio_file, 'rb') as wf:
                    transcript = ""
                    
                    while True:
                        data = wf.readframes(4000)
                        if len(data) == 0:
                            break
                        if recognizer.AcceptWaveform(data):
                            result = json.loads(recognizer.Result())
                            text = result.get("text", "")
                            transcript += text
                    
                    # 获取最后一部分
                    result = json.loads(recognizer.FinalResult())
                    final_text = result.get("text", "")
                    transcript += final_text
            except Exception as e:
                error_msg = f"处理音频文件时出错: {str(e)}"
                print(error_msg)
                import traceback
                traceback.print_exc()
                self.update_status_signal.emit(error_msg)
                return
            
            # 处理AI回答
            print(f"转录结果: '{transcript}'")
            print(f"转录结果长度: {len(transcript)}")
            print(f"转录结果是否为空: {not transcript.strip()}")
            
            # 显示转录结果
            if transcript.strip():
                self.update_transcript_signal.emit(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {transcript}")
                print(f"调用AI回答，文本长度: {len(transcript)}")
                self.get_ai_answer(transcript)
            else:
                print("识别结果为空，不调用AI")
                # 即使识别结果为空，也显示到界面
                self.update_transcript_signal.emit(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] (未识别到语音)")
            
            self.update_status_signal.emit("转录完成")
            print("转录完成")
        except Exception as e:
            error_msg = f"转录失败: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            self.update_status_signal.emit(error_msg)
    
    def get_ai_answer(self, transcript):
        try:
            print(f"开始生成AI回答，问题: '{transcript}'")
            self.update_status_signal.emit("正在生成回答...")
            
            # 从知识库检索相关信息
            context = ""
            if hasattr(self, 'vectorizer') and hasattr(self, 'knowledge_chunks') and self.knowledge_chunks:
                print("开始RAG检索...")
                print(f"知识库大小: {len(self.knowledge_chunks)} 个文本块")
                # 使用TF-IDF向量化器转换查询
                query_vector = self.vectorizer.transform([transcript])
                if hasattr(self, 'knn'):
                    distances, indices = self.knn.kneighbors(query_vector)
                    
                    print(f"检索到的相关索引: {indices[0]}")
                    print(f"检索到的距离: {distances[0]}")
                    
                    # 设置相似度阈值（距离越小，相似度越高）
                    similarity_threshold = 0.7  # 距离阈值，小于此值的认为相似度高
                    
                    for i, (idx, distance) in enumerate(zip(indices[0], distances[0])):
                        if idx < len(self.knowledge_chunks):
                            # 检查相似度
                            if distance < similarity_threshold:
                                chunk = self.knowledge_chunks[idx]
                                context += f"{chunk}\n\n"
                                print(f"添加相关文本块 {i+1}，长度: {len(chunk)}，距离: {distance:.4f}")
                            else:
                                print(f"跳过相似度低的文本块 {i+1}，距离: {distance:.4f}")
                print(f"RAG检索完成，上下文长度: {len(context)}")
            else:
                print("知识库为空或未初始化，跳过RAG检索")
            
            # 构建prompt，要求直接回答问题，不要思考过程
            prompt = f"请直接回答以下问题，不要有任何思考过程，回答要简洁明了：\n\n上下文：{context}\n\n问题：{transcript}\n\n回答："
            
            # 检查OpenAI API Key
            if not self.config.get("openai_api_key"):
                self.update_status_signal.emit("未设置OpenAI API Key")
                # 即使没有API Key，也显示一个默认回答
                self.update_answer_signal.emit(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] 请在配置中设置OpenAI API Key以获取智能回答\n")
                return
            
            # 调用OpenAI API
            from openai import OpenAI
            
            # 检查base_url是否为modelscope
            base_url = self.config.get("openai_base_url", "https://aihubmix.com/v1")
            model_name = self.config.get("openai_model", "coding-minimax-m2.7-free")
            
            print(f"使用OpenAI模型: {model_name}")
            print(f"使用的base_url: {base_url}")
            
            # 检查模型名称是否有效
            if not model_name:
                self.update_status_signal.emit("未设置OpenAI模型名称")
                self.update_answer_signal.emit(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] 请在配置中设置OpenAI模型名称\n")
                return
            
            client = OpenAI(
                api_key=self.config.get("openai_api_key"),
                base_url=base_url

            )
            
            try:
                # 尝试使用completions API
                response = client.completions.create(
                    model=model_name,
                    prompt=prompt,
                    max_tokens=200,  # 减少最大tokens，加快响应
                    temperature=0.1,  # 降低temperature，使回答更直接
                    top_p=0.9,  # 增加top_p，使回答更集中
                    frequency_penalty=0.0,  # 避免重复
                    presence_penalty=0.0
                )
                # 提取回答
                answer = response.choices[0].text.strip()
                print(f"completions API返回: {response}")
                print(f"生成的回答: '{answer}'")
            except Exception as e:
                # 如果completions API失败，尝试使用chat completions API
                print(f"completions API失败，尝试使用chat completions API: {e}")
                # 构建chat messages
                # 修改system message，要求直接回答问题
                messages = [
                    {"role": "system", "content": "你是一个智能助手，直接回答问题，不要有任何思考过程，回答要简洁明了。"},
                    {"role": "user", "content": f"请根据以下上下文回答问题：\n\n上下文：{context}\n\n问题：{transcript}"}
                ]
                response = client.chat.completions.create(
                        model=model_name,
                        messages=messages,
                        max_tokens=200,  # 减少最大tokens，加快响应
                        temperature=0.1,  # 降低temperature，使回答更直接
                        top_p=0.9,  # 增加top_p，使回答更集中
                        frequency_penalty=0.0,  # 避免重复
                        presence_penalty=0.0
                    )
                # 提取回答
                answer = response.choices[0].message.content.strip()
            
            print(f"生成的回答: '{answer}'")
            
            # 显示回答
            self.update_answer_signal.emit(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {answer}\n")
            
            self.update_status_signal.emit("回答生成完成")
        except Exception as e:
            error_msg = f"回答生成失败: {str(e)}"
            print(error_msg)
            self.update_status_signal.emit(error_msg)
            # 即使出错，也显示错误信息
            self.update_answer_signal.emit(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] 回答生成失败: {str(e)}\n")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = VoiceQA()
    window.show()
    sys.exit(app.exec_())
